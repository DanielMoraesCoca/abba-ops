"""Render da minuta final em DOCX — o documento profissional que o advogado edita.

Determinístico (sem LLM): recebe a MinutaFinal (produzida pela Crew de Redação e
aprovada no gate humano) + a trilha de auditoria e monta um .docx no padrão da
casa, com o rodapé obrigatório de revisão/assinatura do advogado nomeado.

O corpo vem em markdown leve (linhas '# '/'## ' viram títulos; resto, parágrafos).
PDF é conversão posterior (LibreOffice/serviço) — o DOCX é o entregável editável.
"""

from __future__ import annotations

from typing import Optional

from docx import Document
from docx.shared import Pt

from patrimonio_flow.schemas import MinutaFinal


def _add_corpo_markdown(doc: "Document", corpo: str) -> None:
    for linha in corpo.splitlines():
        s = linha.rstrip()
        if not s:
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style=None).style = doc.styles["List Bullet"]
        else:
            doc.add_paragraph(s)


def minuta_para_docx(
    minuta: MinutaFinal,
    caminho: str,
    *,
    trilha: str = "",
    preparado_para: str = "",
    titulo: str = "Minuta de Planejamento Patrimonial",
) -> str:
    """Monta o .docx e salva em `caminho`. Retorna o caminho salvo."""
    doc = Document()

    doc.add_heading(titulo, level=0)
    if preparado_para:
        p = doc.add_paragraph()
        p.add_run("Preparado para: ").bold = True
        p.add_run(preparado_para)

    doc.add_heading("Sumário executivo", level=1)
    doc.add_paragraph(minuta.sumario_executivo)

    _add_corpo_markdown(doc, minuta.corpo_markdown)

    if minuta.limites_da_analise:
        doc.add_heading("Limites da análise", level=1)
        for lim in minuta.limites_da_analise:
            doc.add_paragraph(lim, style="List Bullet")

    if minuta.fontes_doc_ids:
        doc.add_heading("Fontes", level=1)
        for fonte in minuta.fontes_doc_ids:
            doc.add_paragraph(fonte, style="List Bullet")

    # rodapé obrigatório (revisão/assinatura do advogado) — sempre presente
    doc.add_paragraph()
    rod = doc.add_paragraph()
    rod.add_run(minuta.rodape_obrigatorio).italic = True

    if trilha:
        t = doc.add_paragraph()
        run = t.add_run(trilha.strip())
        run.italic = True
        run.font.size = Pt(8)

    doc.save(caminho)
    return caminho


def trilha_de_auditoria(versao_corpus: str, n_chunks: int, custo_usd: float) -> str:
    """A linha de trilha padronizada (também usada no render markdown do Flow)."""
    return (f"Trilha: corpus {versao_corpus} · {n_chunks} chunks consultados · "
            f"custo acumulado US$ {custo_usd:.2f}")
